#!/usr/bin/env python
"""One-page-per-option Word brief for PI decision on chemogenetics."""
from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, Rectangle
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve()
OUTDIR = HERE.parents[2] / "outputs"
FIGDIR = OUTDIR / "_pi_word_figs"
FIGDIR.mkdir(parents=True, exist_ok=True)

NAVY = (31 / 255, 59 / 255, 99 / 255)
MORPH = "#C0504D"
WATER = "#5B8FB0"
REST = "#C5B8D9"
TEST = "#1E6F3C"
DCZ = "#1E6F3C"
GREY = "#5A5A5A"
LOCK = "#9C6300"

N_PER = 12
ATTRITION = 1.25
N_GROUPS = 5
N_USE = N_PER * N_GROUPS
N_ORDER = int(round(N_USE * ATTRITION))


def _no_ax(ax):
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")


def _seg(ax, x0, x1, y, h, color, label, fs=8, tc="white"):
    ax.add_patch(Rectangle((x0, y), x1 - x0, h, facecolor=color, edgecolor="white", lw=1.2))
    ax.text((x0 + x1) / 2, y + h / 2, label, ha="center", va="center",
            fontsize=fs, color=tc, fontweight="bold", wrap=True)


def _mark(ax, x, y, h, text, color="#B31B1B"):
    ax.plot([x, x], [y + h, y + h + 0.08], color=color, lw=1.6)
    ax.text(x, y + h + 0.11, text, ha="center", va="bottom", fontsize=7.5,
            color=color, fontweight="bold")


def save(fig, name):
    p = FIGDIR / name
    fig.savefig(p, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return p


def fig_groups(name):
    """Five-group schematic used on every option page."""
    fig, ax = plt.subplots(figsize=(10.2, 1.15))
    _no_ax(ax)
    labels = [
        ("1  Active\nhM4Di  DCZ", "#1E6F3C", "white"),
        ("2  Active\nhM4Di  veh", "#D8ECD8", NAVY),
        ("3  Active\nmCherry  DCZ", "#5A5A5A", "white"),
        ("4  Passive\nhM4Di  DCZ", "#2E6DA4", "white"),
        ("5  Passive\nhM4Di  veh", "#D9E4F0", NAVY),
    ]
    w, gap = 0.175, 0.02
    x = 0.02
    for lab, fill, tc in labels:
        ax.add_patch(FancyBboxPatch((x, 0.12), w, 0.76, boxstyle="round,pad=0.015,rounding_size=0.04",
                                    facecolor=fill, edgecolor="white", lw=0))
        ax.text(x + w / 2, 0.5, lab, ha="center", va="center", fontsize=8,
                color=tc, fontweight="bold", linespacing=1.25)
        x += w + gap
    return save(fig, name)


def fig_A():
    fig, ax = plt.subplots(figsize=(10.2, 1.55))
    _no_ax(ax)
    y, h = 0.28, 0.38
    # 33 day scale
    def x(d):
        return 0.02 + 0.96 * d / 33
    _seg(ax, x(0), x(10), y, h, WATER, "Pre / During", 8, "white")
    _seg(ax, x(10), x(13), y, h, MORPH, "Post", 8)
    _seg(ax, x(13), x(18), y, h, WATER, "Wd / Re-exp", 8, "white")
    _seg(ax, x(18), x(32), y, h, REST, "home cage rest", 8, NAVY)
    _seg(ax, x(32), x(33), y, h, TEST, "TEST", 8)
    _mark(ax, x(12), y, h, "4-OHT")
    _mark(ax, x(33), y, h, "DCZ  ·  pump OFF", "#1E6F3C")
    ax.text(0.5, 0.08, "calendar day 1                                        12                         18                    26 receptor ready           33",
            ha="center", fontsize=7, color=GREY)
    return save(fig, "optA.png")


def fig_B():
    fig, ax = plt.subplots(figsize=(10.2, 1.55))
    _no_ax(ax)
    y, h = 0.28, 0.38
    def x(d):
        return 0.02 + 0.96 * d / 33
    _seg(ax, x(0), x(10), y, h, WATER, "Pre / During", 8, "white")
    _seg(ax, x(10), x(13), y, h, MORPH, "Post", 8)
    _seg(ax, x(13), x(18), y, h, WATER, "Wd / Re-exp", 8, "white")
    _seg(ax, x(18), x(32), y, h, REST, "home cage rest", 8, NAVY)
    _seg(ax, x(32), x(33), y, h, WATER, "water PR", 8, "white")
    _mark(ax, x(12), y, h, "4-OHT")
    _mark(ax, x(33), y, h, "DCZ  ·  water ON", "#5B8FB0")
    ax.text(0.5, 0.08, "same calendar as A   ·   day 33 is your Withdrawal task, not a relapse test",
            ha="center", fontsize=7, color=GREY)
    return save(fig, "optB.png")


def fig_C():
    fig, ax = plt.subplots(figsize=(10.2, 1.55))
    _no_ax(ax)
    y, h = 0.28, 0.38
    def x(d):
        return 0.02 + 0.96 * d / 26
    _seg(ax, x(0), x(10), y, h, WATER, "Pre / During", 8, "white")
    _seg(ax, x(10), x(13), y, h, MORPH, "Post", 8)
    _seg(ax, x(13), x(26), y, h, MORPH, "still morphine  (stretched Post)", 8)
    _mark(ax, x(12), y, h, "4-OHT")
    _mark(ax, x(26), y, h, "DCZ on morphine", "#1E6F3C")
    ax.text(0.5, 0.08, "first cycle only   ·   extra ~2 weeks of morphine   ·   not the protocol behind the TRAP screen",
            ha="center", fontsize=7, color=GREY)
    return save(fig, "optC.png")


def fig_F():
    fig, ax = plt.subplots(figsize=(10.2, 2.15))
    _no_ax(ax)
    def x(d):
        return 0.02 + 0.96 * d / 44
    # run 1
    y1, h = 0.55, 0.28
    _seg(ax, x(0), x(10), y1, h, WATER, "R1 Pre/During", 7.5, "white")
    _seg(ax, x(10), x(13), y1, h, MORPH, "R1 Post", 7.5)
    _seg(ax, x(13), x(18), y1, h, WATER, "R1 Wd/Re", 7.5, "white")
    _seg(ax, x(18), x(25), y1, h, REST, "rest", 7.5, NAVY)
    _mark(ax, x(12), y1, h, "4-OHT")
    # run 2
    y2 = 0.14
    _seg(ax, x(25), x(35), y2, h, WATER, "R2 Pre / During   NO DCZ", 7.5, "white")
    _seg(ax, x(35), x(38), y2, h, DCZ, "R2 Post", 7.5)
    _seg(ax, x(38), x(41), y2, h, WATER, "R2 Wd", 7.5, "white")
    _seg(ax, x(41), x(44), y2, h, MORPH, "R2 Re-exp", 7.5)
    ax.annotate("", xy=(x(25), y2 + h), xytext=(x(25), y1),
                arrowprops=dict(arrowstyle="-", color="#888", lw=0.8))
    ax.text(x(36.5), y2 + h + 0.04, "DCZ on ONE of these three  (circle below)",
            ha="center", fontsize=7.5, color="#1E6F3C", fontweight="bold")
    ax.text(0.5, 0.02, "yoking intact in Run-2 During   ·   do not DCZ During   ·   do not DCZ all three phases in the same mouse",
            ha="center", fontsize=7, color=GREY)
    return save(fig, "optF.png")


def fig_E():
    fig, ax = plt.subplots(figsize=(10.2, 1.55))
    _no_ax(ax)
    y, h = 0.28, 0.38
    def x(d):
        return 0.02 + 0.96 * d / 18
    _seg(ax, x(0), x(10), y, h, WATER, "Pre / During", 8, "white")
    _seg(ax, x(10), x(13), y, h, MORPH, "Post", 8)
    _seg(ax, x(13), x(18), y, h, WATER, "Wd / Re-exp", 8, "white")
    _mark(ax, x(12), y, h, "4-OHT")
    _mark(ax, x(13), y, h, "DCZ possible  (+1 day)", "#1E6F3C")
    ax.text(0.5, 0.08, "TRAP2 × R26-hM4Di  ·  48 h, not 14 days  ·  original 3-day Post  ·  start breeding now",
            ha="center", fontsize=7, color=GREY)
    return save(fig, "optE.png")


# ---- Word helpers ----------------------------------------------------------

def shade(cell, hex_color: str):
    tc = cell._tePr if False else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run(run, size=11, bold=False, color=(31, 59, 99), font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font)


def p(doc, text, size=11, bold=False, color=(31, 59, 99), space_after=4, center=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.05
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return para


def banner(doc, title, subtitle):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    shade(cell, "1F3B63")
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = para.add_run(title + "\n")
    set_run(r, size=18, bold=True, color=(255, 255, 255))
    r2 = para.add_run(subtitle)
    set_run(r2, size=12, bold=False, color=(201, 218, 236))
    for pp in cell.paragraphs:
        pp.paragraph_format.space_before = Pt(6)
        pp.paragraph_format.space_after = Pt(6)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, rows, col_w, header_fill="1F3B63"):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j < 4 else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(str(val))
            if i == 0:
                shade(cell, header_fill)
                set_run(run, size=10, bold=True, color=(255, 255, 255))
            else:
                shade(cell, "FFFFFF" if i % 2 else "EEF3F9")
                set_run(run, size=10, bold=(j == 0), color=(31, 59, 99))
            cell.width = Inches(col_w[j])
    return t


def landscape(doc):
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Inches(11)
    sec.page_height = Inches(8.5)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)
    sec.top_margin = Inches(0.4)
    sec.bottom_margin = Inches(0.4)


def new_page(doc):
    doc.add_page_break()


def option_page(doc, code, title, question, pro, con, fig_path, extra_rows=None, note=""):
    banner(doc, f"Option {code}   ·   {title}", question)
    p(doc, f"PRO  {pro}      ·      CON  {con}", size=11, bold=False, space_after=6)
    doc.add_picture(str(fig_path), width=Inches(10.0))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    last.paragraph_format.space_after = Pt(4)

    p(doc, "Groups  (same for A, B, C, and F)     n = 12 per cell     use 60     order 75 for ORBm",
      size=11, bold=True, space_after=4)
    doc.add_picture(str(FIGDIR / "groups.png"), width=Inches(10.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows = [
        ["#", "Mice", "Virus", "Drug on test days", "n", "Order"],
        ["1", "Active", "hM4Di-mCherry", "DCZ 0.1 mg/kg", "12", "15"],
        ["2", "Active", "hM4Di-mCherry", "vehicle", "12", "15"],
        ["3", "Active", "mCherry", "DCZ 0.1 mg/kg", "12", "15"],
        ["4", "Passive", "hM4Di-mCherry", "DCZ 0.1 mg/kg", "12", "15"],
        ["5", "Passive", "hM4Di-mCherry", "vehicle", "12", "15"],
    ]
    if extra_rows:
        rows.extend(extra_rows)
    add_table(doc, rows, [0.6, 1.3, 2.0, 2.4, 0.7, 0.8])
    if note:
        p(doc, note, size=10, color=(90, 90, 90), space_after=0)


def cover(doc):
    banner(doc, "Chemogenetic test of the Post ensemble",
           "Circle one option. Locked: TRAP2 × Ai14  ·  4-OHT after Post day 12  ·  wait 14 days  ·  DCZ  ·  Active + Passive  ·  ORBm first")
    p(doc, "What is locked vs what you decide", size=13, bold=True, space_after=6)
    add_table(doc, [
        ["Locked", "You decide"],
        ["Tag at first Post (day 12 of the 18-day paradigm)", "Which question this year"],
        ["AAV-DIO-hM4Di into ORBm 3 weeks before tagging", "A / B / C / F  (E is next year)"],
        ["First possible DCZ = day 26  (14 days after 4-OHT)", "If F: which Run-2 phase gets DCZ"],
        ["No DCZ during During  (yoking stays intact)", "Then BMAp, same design"],
        ["One mouse = one seeking/taking test type per question", ""],
    ], [5.0, 5.0])

    p(doc, "", size=6, space_after=8)
    p(doc, "The five options  (one page each)", size=13, bold=True, space_after=6)
    add_table(doc, [
        ["", "Question", "Mice (use / order)", "Pick?"],
        ["A", "After rest: does silencing cut cue-seeking?  (pump OFF)", f"{N_USE} / {N_ORDER}", "☐"],
        ["B", "After rest: does silencing cut water-PR?  (your Withdrawal task)", f"{N_USE} / {N_ORDER}", "☐"],
        ["C", "Stretch first Post: silence while still on morphine", f"{N_USE} / {N_ORDER}", "☐"],
        ["F", "Second 18-day cycle: silence Post or Wd or Re-exp  (not During)", f"{N_USE} / {N_ORDER} per phase", "☐"],
        ["E", "Next year: new line, silence inside the original 3-day Post", "breeding first", "☐"],
    ], [0.7, 5.6, 2.4, 0.8], header_fill="1F3B63")

    p(doc, "", size=4, space_after=6)
    p(doc, "Recommendation:  A this year if the claim is relapse.  "
           "F this year if the claim is “first-Post cells are needed in a second cycle.”  "
           "Start E breeding now if the original 3-day Post still matters.  "
           "Do not DCZ the same mouse on Post and Withdrawal and Reinstatement.",
      size=11, bold=True, space_after=0)


def main() -> None:
    fig_groups("groups.png")
    fa, fb, fc, ff, fe = fig_A(), fig_B(), fig_C(), fig_F(), fig_E()

    doc = Document()
    landscape(doc)
    cover(doc)

    new_page(doc)
    option_page(
        doc, "A", "Late relapse test",
        "Run the 18-day paradigm once. Tag day 12. Rest. One test on day 33. Pump OFF. Nothing delivered.",
        "Receptor ready (+21 d). Standard relapse claim.",
        "New session type. Not your Withdrawal. Not Q1.",
        fa,
        note="1 vs 2 = hypothesis.  1 vs 3 = ligand.  4 vs 5 = Active-only?   "
             "Open field 10 min after DCZ, before the test, is free locomotion control.",
    )

    new_page(doc)
    option_page(
        doc, "B", "Late water-PR  (your Withdrawal task)",
        "Same calendar as A. Day 33 is water PR, pump ON. Matches day 14–16 of your protocol.",
        "Same task as existing TRAP Withdrawal data.",
        "Water is available, so this is weak as “drug seeking.”",
        fb,
        note="Use this only if the claim is “late Withdrawal-like licking,” not relapse.",
    )

    new_page(doc)
    option_page(
        doc, "C", "Stretch Post on the first cycle",
        "Do not rest. Keep morphine until day 26. Then DCZ while the mouse is still taking.",
        "Tests Q1 this year (silence while taking).",
        "Extra ~2 weeks of morphine. Drug history ≠ the TRAP screen.",
        fc,
        note="Yoking is not an issue here if DCZ is after During.  "
             "This is not comparable to current whole-brain TRAP densities.",
    )

    new_page(doc)
    option_page(
        doc, "F", "Second cycle  —  DCZ on ONE later phase",
        "Run 1 = full 18 days, 4-OHT day 12, no DCZ. Rest. Run 2 = full 18 days. "
        "No DCZ in During (yoking intact). DCZ on exactly one of: Post / Withdrawal / Re-exposure.",
        "During of Run 2 is a clean no-DCZ baseline. Tests whether first-Post cells matter later.",
        "Second cycle ≠ first-cycle data. Three phases in one mouse would contaminate each other.",
        ff,
        note="Circle the Run-2 phase for this cohort:   ☐ Post (morphine PR)   "
             "☐ Withdrawal (water PR)   ☐ Re-exposure (morphine PR).   "
             "All 3 days of that phase = DCZ or all 3 = vehicle. No alternating days.  "
             "A second phase = another 75 mice.",
    )

    new_page(doc)
    banner(doc, "Option E   ·   Phase 2  (not this cohort)",
           "TRAP2 × R26-hM4Di + local DCZ. Receptor ready 48 h after 4-OHT. Original 3-day Post can be tested.")
    p(doc, "PRO  True Q1 inside your existing 18-day map.     CON  Breed 4–6 months. Cannot also carry Ai14. Needs cannulae.",
      size=11, space_after=6)
    doc.add_picture(str(fe), width=Inches(10.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p(doc, "Do not wait to decide E. If Q1 still matters, start the cross in parallel with A or F.",
      size=12, bold=True, space_after=8)
    add_table(doc, [
        ["Item", "Detail"],
        ["Line", "TRAP2 (JAX 030323) × R26-LSL-hM4Di (JAX 026219)"],
        ["Why not Ai14", "Both sit at Rosa26"],
        ["Region", "Local DCZ through cannulae in ORBm (transgene is brain-wide)"],
        ["n when ready", "Same 5 groups × 12  →  order 75"],
        ["Start", "Breed now; first mice ~mid 2027"],
    ], [2.2, 7.8])

    out = OUTDIR / "Chemo_PI_decision.docx"
    doc.save(out)
    print(f"[DONE] {out}")


if __name__ == "__main__":
    main()
